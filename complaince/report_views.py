from __future__ import annotations

import io
from datetime import datetime
from typing import Dict, Iterable, List, Optional, Set, Tuple

from django.db.models import Count, Q
from django.http import FileResponse
from django.utils import timezone
from rest_framework import status
from rest_framework.response import Response
from rest_framework.views import APIView

from community.models import CensusYear, Community, CommunityCensusData
from sites.models import SiteCensusData

from .models import ComplianceCalculation


PROGRAM_FIELD = {
    "Paint": "program_paint",
    "Lighting": "program_lights",
    "Solvents": "program_solvents",
    "Pesticides": "program_pesticides",
    "Fertilizers": "program_fertilizers",
}


def _require_superuser(request) -> Optional[Response]:
    user = getattr(request, "user", None)
    if not user or not user.is_authenticated:
        return Response({"detail": "Authentication credentials were not provided."}, status=401)
    if not getattr(user, "is_superuser", False):
        return Response({"detail": "You do not have permission to perform this action."}, status=403)
    return None


def _resolve_census_year(census_year_id, year_value) -> Tuple[Optional[CensusYear], Optional[Response]]:
    if census_year_id not in (None, ""):
        try:
            return CensusYear.objects.get(id=census_year_id), None
        except (CensusYear.DoesNotExist, ValueError, TypeError):
            return None, Response(
                {"error": f"Census year id {census_year_id} not found."}, status=status.HTTP_404_NOT_FOUND
            )
    if year_value not in (None, ""):
        try:
            return CensusYear.objects.get(year=int(year_value)), None
        except (CensusYear.DoesNotExist, ValueError, TypeError):
            return None, Response(
                {"error": f"Census year {year_value!r} not found."}, status=status.HTTP_404_NOT_FOUND
            )
    return None, Response(
        {
            "error": "census_year_id or year is required.",
            "hint": "Use census_year_id (pk) or year (calendar year), e.g. ?year=2000.",
        },
        status=status.HTTP_400_BAD_REQUEST,
    )


def _parse_date(d: str) -> Optional[datetime]:
    if not d:
        return None
    try:
        # Accept yyyy-mm-dd (from HTML date input)
        return datetime.strptime(d, "%Y-%m-%d")
    except ValueError:
        return None


def _site_date_filter_q(filter_type: str, start_date: Optional[datetime], end_date: Optional[datetime]) -> Q:
    """
    filterType in UI: "activated", "deactivated", "all".
    We approximate this using site_start_date/site_end_date.
    """
    if filter_type not in ("activated", "deactivated"):
        return Q()
    if not start_date and not end_date:
        return Q()
    if start_date:
        start = timezone.make_aware(start_date) if timezone.is_naive(start_date) else start_date
    else:
        start = None
    if end_date:
        # include whole day
        end = timezone.make_aware(end_date) if timezone.is_naive(end_date) else end_date
        end = end.replace(hour=23, minute=59, second=59, microsecond=999999)
    else:
        end = None

    if filter_type == "activated":
        q = Q(site_start_date__isnull=False)
        if start:
            q &= Q(site_start_date__gte=start)
        if end:
            q &= Q(site_start_date__lte=end)
        return q

    q = Q(site_end_date__isnull=False)
    if start:
        q &= Q(site_end_date__gte=start)
    if end:
        q &= Q(site_end_date__lte=end)
    return q


def _normalize_programs(programs: Iterable[str]) -> List[str]:
    out = []
    for p in programs or []:
        if p in PROGRAM_FIELD:
            out.append(p)
    return sorted(set(out))


def _community_ids_from_selection(
    census_year: CensusYear,
    selected_names: Iterable[str],
) -> Set[str]:
    names = [str(n).strip() for n in (selected_names or []) if str(n).strip()]
    if not names:
        return set()
    # UI passes municipality.name
    return set(
        Community.objects.filter(name__in=names).values_list("id", flat=True)
    )


class ReportConfigView(APIView):
    """
    GET /api/compliance/reports/config/?year=2000
    Returns report types + program counts + municipality list (dedupe-ready) with site counts.
    """

    def get(self, request):
        census_year, err = _resolve_census_year(
            request.query_params.get("census_year_id"),
            request.query_params.get("year"),
        )
        if err:
            return err

        # Programs: active site counts per program (quick grouped counts)
        programs = []
        base_qs = SiteCensusData.objects.filter(census_year=census_year, is_active=True)
        for prog, field in PROGRAM_FIELD.items():
            cnt = base_qs.filter(**{field: True}).count()
            programs.append({"program": prog, "active_sites": cnt})

        # Municipalities: use CommunityCensusData to constrain to year
        communities = (
            CommunityCensusData.objects.filter(census_year=census_year)
            .select_related("community")
            .order_by("community__name")
        )
        community_ids = [c.community_id for c in communities]

        # One grouped query for "active sites in community" (any program)
        active_any = (
            SiteCensusData.objects.filter(census_year=census_year, is_active=True, community_id__in=community_ids)
            .values("community_id")
            .annotate(c=Count("id"))
        )
        active_any_map = {str(r["community_id"]): int(r["c"]) for r in active_any}

        municipalities = []
        for cd in communities:
            municipalities.append(
                {
                    "id": str(cd.community_id),
                    "name": cd.community.name,
                    "population": cd.population,
                    "tier": cd.tier,
                    "active_sites": active_any_map.get(str(cd.community_id), 0),
                }
            )

        report_types = [
            {"value": "compliance-summary", "label": "Compliance Summary Report"},
            {"value": "site-inventory", "label": "Site Inventory Report"},
            {"value": "shortfall-analysis", "label": "Shortfall Analysis Report"},
            {"value": "reallocation-report", "label": "Reallocation Report"},
            {"value": "offset-report", "label": "Offset Report"},
            {"value": "regulatory-compliance", "label": "Regulatory Compliance Report"},
            {"value": "historical-tracking", "label": "Historical Site Changes"},
            {"value": "site-type-breakdown", "label": "Site Type Breakdown Analysis"},
        ]

        return Response(
            {
                "census_year": {"id": census_year.id, "year": census_year.year},
                "report_types": report_types,
                "programs": programs,
                "municipalities": municipalities,
            }
        )


class ReportPreviewView(APIView):
    """
    POST /api/compliance/reports/preview/
    Body (example):
      {
        "report_type": "compliance-summary",
        "year": 2000,
        "programs": ["Paint","Lighting"],
        "municipalities": ["Toronto","Ottawa"],
        "date_filter": {"startDate":"", "endDate":"", "filterType":"all"},
        "options": {"include_charts": true, "include_details": true}
      }
    """

    def post(self, request):
        data = request.data or {}
        census_year, err = _resolve_census_year(data.get("census_year_id"), data.get("year"))
        if err:
            return err

        report_type = data.get("report_type") or "compliance-summary"
        selected_programs = _normalize_programs(data.get("programs") or [])
        if not selected_programs:
            # Default: all programs (matches UI initial selection behavior)
            selected_programs = sorted(PROGRAM_FIELD.keys())

        selected_municipalities = data.get("municipalities") or []
        selected_ids = _community_ids_from_selection(census_year, selected_municipalities)

        df = data.get("date_filter") or {}
        filter_type = df.get("filterType") or "all"
        start_date = _parse_date(df.get("startDate") or "")
        end_date = _parse_date(df.get("endDate") or "")

        site_q = Q(census_year=census_year, is_active=True)
        if selected_ids:
            site_q &= Q(community_id__in=selected_ids)
        site_q &= _site_date_filter_q(filter_type, start_date, end_date)

        # Active sites counts (any of selected programs)
        program_q = Q()
        for p in selected_programs:
            program_q |= Q(**{PROGRAM_FIELD[p]: True})
        sites_qs = SiteCensusData.objects.filter(site_q & program_q).select_related("community", "census_year")
        total_sites = sites_qs.count()

        # Municipalities base set (use CommunityCensusData)
        muni_qs = CommunityCensusData.objects.filter(census_year=census_year).select_related("community")
        if selected_ids:
            muni_qs = muni_qs.filter(community_id__in=selected_ids)
        muni_rows = list(muni_qs)
        muni_ids = [m.community_id for m in muni_rows]

        # Compliance aggregation by community across selected programs
        calcs = (
            ComplianceCalculation.objects.filter(census_year=census_year, program__in=selected_programs, community_id__in=muni_ids)
            .values("community_id")
        )
        # Aggregate in python (required/actual stored on calc)
        by_comm: Dict[str, Dict[str, int]] = {}
        for cc in ComplianceCalculation.objects.filter(
            census_year=census_year, program__in=selected_programs, community_id__in=muni_ids
        ).only("community_id", "required_sites", "actual_sites"):
            key = str(cc.community_id)
            d = by_comm.setdefault(key, {"required": 0, "actual": 0})
            d["required"] += int(cc.required_sites or 0)
            d["actual"] += int(cc.actual_sites or 0)

        total_shortfalls = 0
        total_excesses = 0
        compliant_munis = 0
        for mid in muni_ids:
            k = str(mid)
            req = by_comm.get(k, {}).get("required", 0)
            act = by_comm.get(k, {}).get("actual", 0)
            shortfall = max(0, req - act)
            excess = max(0, act - req)
            total_shortfalls += shortfall
            total_excesses += excess
            if shortfall == 0:
                compliant_munis += 1

        # Program breakdown: sites + municipalities served
        program_breakdown = []
        for p in selected_programs:
            field = PROGRAM_FIELD[p]
            p_sites = SiteCensusData.objects.filter(site_q, **{field: True})
            program_breakdown.append(
                {
                    "program": p,
                    "active_sites": p_sites.count(),
                    "municipalities_served": p_sites.values("community_id").distinct().count(),
                    "status": "Active",
                }
            )

        # Municipality summary
        # Precompute active sites per municipality (selected programs)
        sites_by_muni = (
            sites_qs.values("community_id").annotate(c=Count("id"))
        )
        sites_by_muni_map = {str(r["community_id"]): int(r["c"]) for r in sites_by_muni}
        # Precompute programs served per municipality
        programs_served_map: Dict[str, Set[str]] = {str(mid): set() for mid in muni_ids}
        for p in selected_programs:
            field = PROGRAM_FIELD[p]
            for cid in SiteCensusData.objects.filter(site_q, **{field: True}).values_list("community_id", flat=True).distinct():
                programs_served_map.setdefault(str(cid), set()).add(p)

        municipality_summary = []
        for m in muni_rows:
            municipality_summary.append(
                {
                    "municipality": m.community.name,
                    "population": m.population,
                    "active_sites": sites_by_muni_map.get(str(m.community_id), 0),
                    "programs_served": sorted(programs_served_map.get(str(m.community_id), set())),
                    "tier": m.tier,
                }
            )

        compliance_data = {
            "totalSites": total_sites,
            "compliantMunicipalities": compliant_munis,
            "totalMunicipalities": len(muni_ids),
            "shortfalls": total_shortfalls,
            "excesses": total_excesses,
        }

        return Response(
            {
                "census_year": {"id": census_year.id, "year": census_year.year},
                "report_type": report_type,
                "filters": {
                    "programs": selected_programs,
                    "municipalities": selected_municipalities,
                    "date_filter": {"filterType": filter_type, "startDate": df.get("startDate") or "", "endDate": df.get("endDate") or ""},
                },
                "complianceData": compliance_data,
                "programBreakdown": program_breakdown,
                "municipalitySummary": municipality_summary,
            }
        )


class ReportExportView(APIView):
    """
    POST /api/compliance/reports/export/
    Generates and returns a file response (xlsx/docx/pdf).
    """

    def post(self, request):
        data = request.data or {}
        export_format = (data.get("format") or "excel").lower()
        if export_format not in ("excel", "word", "pdf"):
            return Response({"error": "format must be one of excel|word|pdf"}, status=400)

        # Reuse preview payload so exports match UI expectations
        preview = ReportPreviewView().post(request)
        if preview.status_code != 200:
            return preview
        payload = preview.data

        title = payload.get("report_type") or "report"
        year = (payload.get("census_year") or {}).get("year")
        safe_name = f"{title}-{year}" if year else f"{title}"

        if export_format == "excel":
            from openpyxl import Workbook
            from openpyxl.utils import get_column_letter

            wb = Workbook()
            ws = wb.active
            ws.title = "Executive Summary"

            cd = payload.get("complianceData") or {}
            ws.append(["Report Type", payload.get("report_type")])
            ws.append(["Census Year", year])
            ws.append([])
            ws.append(["Total Sites", cd.get("totalSites", 0)])
            ws.append(["Compliance Rate (%)", _rate(cd.get("compliantMunicipalities", 0), cd.get("totalMunicipalities", 0))])
            ws.append(["Shortfalls", cd.get("shortfalls", 0)])
            ws.append(["Excesses", cd.get("excesses", 0)])

            ws2 = wb.create_sheet("Program Breakdown")
            ws2.append(["Program", "Active Sites", "Municipalities Served", "Status"])
            for r in payload.get("programBreakdown") or []:
                ws2.append([r.get("program"), r.get("active_sites"), r.get("municipalities_served"), r.get("status")])

            ws3 = wb.create_sheet("Municipality Summary")
            ws3.append(["Municipality", "Population", "Active Sites", "Programs Served", "Tier"])
            for r in payload.get("municipalitySummary") or []:
                ws3.append([r.get("municipality"), r.get("population"), r.get("active_sites"), ", ".join(r.get("programs_served") or []), r.get("tier")])

            for sheet in (ws, ws2, ws3):
                for col in range(1, sheet.max_column + 1):
                    sheet.column_dimensions[get_column_letter(col)].width = 22

            buf = io.BytesIO()
            wb.save(buf)
            buf.seek(0)
            return FileResponse(
                buf,
                as_attachment=True,
                filename=f"{safe_name}.xlsx",
                content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )

        if export_format == "word":
            try:
                from docx import Document
            except Exception:
                return Response(
                    {"error": "Word export requires python-docx. Add it to requirements and deploy."},
                    status=501,
                )

            doc = Document()
            doc.add_heading("Compliance Report", level=1)
            doc.add_paragraph(f"Report Type: {payload.get('report_type')}")
            doc.add_paragraph(f"Census Year: {year}")

            cd = payload.get("complianceData") or {}
            doc.add_heading("Executive Summary", level=2)
            doc.add_paragraph(f"Total Sites: {cd.get('totalSites', 0)}")
            doc.add_paragraph(
                f"Compliance Rate: {_rate(cd.get('compliantMunicipalities', 0), cd.get('totalMunicipalities', 0))}%"
            )
            doc.add_paragraph(f"Shortfalls: {cd.get('shortfalls', 0)}")
            doc.add_paragraph(f"Excesses: {cd.get('excesses', 0)}")

            doc.add_heading("Program Breakdown", level=2)
            table = doc.add_table(rows=1, cols=4)
            hdr = table.rows[0].cells
            hdr[0].text = "Program"
            hdr[1].text = "Active Sites"
            hdr[2].text = "Municipalities Served"
            hdr[3].text = "Status"
            for r in payload.get("programBreakdown") or []:
                row = table.add_row().cells
                row[0].text = str(r.get("program") or "")
                row[1].text = str(r.get("active_sites") or 0)
                row[2].text = str(r.get("municipalities_served") or 0)
                row[3].text = str(r.get("status") or "")

            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            return FileResponse(
                buf,
                as_attachment=True,
                filename=f"{safe_name}.docx",
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

        # PDF
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
        except Exception:
            return Response(
                {"error": "PDF export requires reportlab. Add it to requirements and deploy."},
                status=501,
            )

        buf = io.BytesIO()
        c = canvas.Canvas(buf, pagesize=letter)
        width, height = letter
        y = height - 50
        c.setFont("Helvetica-Bold", 14)
        c.drawString(50, y, "Compliance Report")
        y -= 22
        c.setFont("Helvetica", 10)
        c.drawString(50, y, f"Report Type: {payload.get('report_type')}")
        y -= 14
        c.drawString(50, y, f"Census Year: {year}")
        y -= 20
        cd = payload.get("complianceData") or {}
        c.setFont("Helvetica-Bold", 12)
        c.drawString(50, y, "Executive Summary")
        y -= 16
        c.setFont("Helvetica", 10)
        lines = [
            f"Total Sites: {cd.get('totalSites', 0)}",
            f"Compliance Rate: {_rate(cd.get('compliantMunicipalities', 0), cd.get('totalMunicipalities', 0))}%",
            f"Shortfalls: {cd.get('shortfalls', 0)}",
            f"Excesses: {cd.get('excesses', 0)}",
        ]
        for line in lines:
            c.drawString(60, y, line)
            y -= 14
        c.showPage()
        c.save()
        buf.seek(0)
        return FileResponse(buf, as_attachment=True, filename=f"{safe_name}.pdf", content_type="application/pdf")


def _rate(numer: int, denom: int) -> int:
    try:
        numer = int(numer or 0)
        denom = int(denom or 0)
    except Exception:
        return 0
    if denom <= 0:
        return 0
    return int(round((numer / denom) * 100))

